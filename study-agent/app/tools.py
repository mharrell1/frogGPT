# Copyright 2026 Google LLC
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     https://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Custom tools for the Study Agent."""

import os


def read_file_content(file_path: str) -> dict:
    """Read the text content of a local file (TXT, MD, CSV, or basic DOCX/PDF via text extraction).

    Args:
        file_path: Absolute or relative path to the file to read.

    Returns:
        A dict with 'status' and 'content' keys. Content is the extracted text.
    """
    if not os.path.exists(file_path):
        return {
            "status": "error",
            "content": f"File not found: {file_path}. Please check the path and try again.",
        }

    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext in (".txt", ".md", ".csv", ".rst"):
            with open(file_path, encoding="utf-8") as f:
                content = f.read()
            return {"status": "success", "content": content, "format": ext}

        elif ext == ".pdf":
            try:
                import pypdf  # type: ignore

                reader = pypdf.PdfReader(file_path)
                pages = [page.extract_text() or "" for page in reader.pages]
                content = "\n\n".join(pages)
                return {"status": "success", "content": content, "format": "pdf"}
            except ImportError:
                return {
                    "status": "error",
                    "content": (
                        "pypdf is not installed. Please ask the user to paste the text content "
                        "directly, or install pypdf: uv add pypdf"
                    ),
                }

        elif ext in (".docx",):
            try:
                import docx  # type: ignore

                doc = docx.Document(file_path)
                content = "\n".join(
                    para.text for para in doc.paragraphs if para.text.strip()
                )
                return {"status": "success", "content": content, "format": "docx"}
            except ImportError:
                return {
                    "status": "error",
                    "content": (
                        "python-docx is not installed. Please ask the user to paste the text "
                        "content directly, or install it: uv add python-docx"
                    ),
                }

        else:
            # Fallback: try reading as plain text
            with open(file_path, encoding="utf-8", errors="replace") as f:
                content = f.read()
            return {
                "status": "success",
                "content": content,
                "format": "unknown (read as text)",
            }

    except Exception as e:
        return {"status": "error", "content": f"Failed to read file: {e}"}


def format_flashcards_markdown(flashcard_deck_json: str) -> dict:
    """Convert a FlashcardDeck JSON string into a human-readable Markdown document.

    Args:
        flashcard_deck_json: JSON string representation of a FlashcardDeck object.

    Returns:
        A dict with 'status' and 'markdown' keys containing the formatted output.
    """
    import json

    try:
        deck = json.loads(flashcard_deck_json)
        lines = [
            f"# 📇 {deck.get('title', 'Flashcard Deck')}",
            f"\n_{deck.get('description', '')}_",
            f"\n---\n\n**{len(deck.get('cards', []))} Cards**\n",
        ]
        for i, card in enumerate(deck.get("cards", []), 1):
            topic = card.get("topic", "")
            topic_str = f" _{topic}_" if topic else ""
            lines.append(f"### Card {i}{topic_str}")
            lines.append(f"\n**Q:** {card.get('question', '')}\n")
            lines.append(f"**A:** {card.get('answer', '')}\n")
            lines.append("---\n")
        return {"status": "success", "markdown": "\n".join(lines)}
    except Exception as e:
        return {"status": "error", "markdown": f"Failed to format flashcards: {e}"}


def format_study_guide_markdown(study_guide_json: str) -> dict:
    """Convert a StudyGuide JSON string into a human-readable Markdown document.

    Args:
        study_guide_json: JSON string representation of a StudyGuide object.

    Returns:
        A dict with 'status' and 'markdown' keys containing the formatted output.
    """
    import json

    try:
        guide = json.loads(study_guide_json)
        lines = [
            f"# 📚 {guide.get('title', 'Study Guide')}",
            f"\n> {guide.get('overview', '')}\n",
        ]
        for section in guide.get("sections", []):
            lines.append(f"\n## {section.get('heading', 'Section')}")
            lines.append(f"\n{section.get('summary', '')}\n")
            concepts = section.get("key_concepts", [])
            if concepts:
                lines.append("\n### Key Concepts\n")
                for concept in concepts:
                    lines.append(f"**{concept.get('term', '')}**")
                    lines.append(f": {concept.get('definition', '')}")
                    if concept.get("example"):
                        lines.append(f"\n  > 💡 *Example:* {concept['example']}\n")
                    else:
                        lines.append("")
            bullets = section.get("bullet_points", [])
            if bullets:
                lines.append("\n### Key Points\n")
                for bullet in bullets:
                    lines.append(f"- {bullet}")
                lines.append("")
        summary = guide.get("summary", "")
        if summary:
            lines.append(f"\n---\n\n## 📌 Summary\n\n{summary}")
        return {"status": "success", "markdown": "\n".join(lines)}
    except Exception as e:
        return {"status": "error", "markdown": f"Failed to format study guide: {e}"}


def format_quiz_markdown(quiz_json: str) -> dict:
    """Convert a Quiz JSON string into a formatted Markdown quiz with answer key.

    Args:
        quiz_json: JSON string representation of a Quiz object.

    Returns:
        A dict with 'status' and 'markdown' keys containing the formatted output.
    """
    import json

    try:
        quiz = json.loads(quiz_json)
        difficulty = quiz.get("difficulty", "medium")
        lines = [
            f"# 📝 {quiz.get('title', 'Practice Quiz')}",
            f"\n**Topic:** {quiz.get('topic', '')} | **Difficulty:** {difficulty.capitalize()}",
            f"**Questions:** {len(quiz.get('questions', []))}\n",
            "\n---\n",
        ]
        answer_key = []
        for i, q in enumerate(quiz.get("questions", []), 1):
            lines.append(f"**{i}.** {q.get('question', '')}\n")
            for opt in q.get("options", []):
                lines.append(f"   {opt}")
            lines.append("")
            answer_key.append(
                f"**{i}.** {q.get('correct_answer', '')} — {q.get('explanation', '')}"
            )

        lines.append("\n---\n\n## ✅ Answer Key\n")
        lines.extend(answer_key)
        return {"status": "success", "markdown": "\n".join(lines)}
    except Exception as e:
        return {"status": "error", "markdown": f"Failed to format quiz: {e}"}


def format_practice_test_markdown(test_json: str) -> dict:
    """Convert a PracticeTest JSON string into a formatted Markdown test with answer key.

    Args:
        test_json: JSON string representation of a PracticeTest object.

    Returns:
        A dict with 'status' and 'markdown' keys containing the formatted output.
    """
    import json

    try:
        test = json.loads(test_json)
        q_num = 1
        lines = [
            f"# 🧪 {test.get('title', 'Practice Test')}",
            f"\n**Topic:** {test.get('topic', '')}",
            f"**Total Points:** {test.get('total_points', '?')} | "
            f"**Time Limit:** {test.get('time_limit_minutes', '?')} minutes\n",
            f"\n> {test.get('instructions', '')}\n",
            "\n---\n",
        ]
        answer_key = []

        # True/False
        tf_qs = test.get("true_false", [])
        if tf_qs:
            lines.append("## Section A — True or False\n")
            for q in tf_qs:
                pts = q.get("points", 1)
                lines.append(
                    f"**{q_num}.** ({pts} pt{'s' if pts != 1 else ''}) {q.get('statement', '')}"
                )
                lines.append("   ☐ True   ☐ False\n")
                answer_key.append(
                    f"**{q_num}.** {q.get('correct_answer', '')} — {q.get('explanation', '')}"
                )
                q_num += 1

        # Multiple Choice
        mc_qs = test.get("multiple_choice", [])
        if mc_qs:
            lines.append("## Section B — Multiple Choice\n")
            for q in mc_qs:
                pts = q.get("points", 1)
                lines.append(
                    f"**{q_num}.** ({pts} pt{'s' if pts != 1 else ''}) {q.get('question', '')}\n"
                )
                for opt in q.get("options", []):
                    lines.append(f"   {opt}")
                lines.append("")
                answer_key.append(
                    f"**{q_num}.** {q.get('correct_answer', '')} — {q.get('explanation', '')}"
                )
                q_num += 1

        # Short Answer
        sa_qs = test.get("short_answer", [])
        if sa_qs:
            lines.append("## Section C — Short Answer\n")
            for q in sa_qs:
                pts = q.get("points", 3)
                lines.append(f"**{q_num}.** ({pts} pts) {q.get('question', '')}\n")
                lines.append(
                    "   _Answer:_ \\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\n"
                )
                kp = q.get("key_points", [])
                answer_key.append(
                    f"**{q_num}.** Sample: {q.get('sample_answer', '')} "
                    f"| Key points: {', '.join(kp)}"
                )
                q_num += 1

        lines.append("\n---\n\n## ✅ Answer Key (Instructor Copy)\n")
        lines.extend(answer_key)
        return {"status": "success", "markdown": "\n".join(lines)}
    except Exception as e:
        return {"status": "error", "markdown": f"Failed to format practice test: {e}"}


def format_explanation_markdown(explanation_json: str) -> dict:
    """Convert a TopicExplanation JSON string into a human-readable Markdown document.

    Args:
        explanation_json: JSON string representation of a TopicExplanation object.

    Returns:
        A dict with 'status' and 'markdown' keys containing the formatted output.
    """
    import json

    try:
        exp = json.loads(explanation_json)
        lines = [
            f"# 💡 Explanation: {exp.get('topic', 'Topic')}",
            f"\n### 📌 High-Level Summary\n{exp.get('summary', '')}\n",
            f"### 👶 Simplified Explanation\n{exp.get('simple_explanation', '')}\n",
            f"### 🎭 Analogy / Metaphor\n> {exp.get('analogy', '')}\n",
            "\n### 🗝️ Key Takeaways",
        ]
        for takeaway in exp.get("key_takeaways", []):
            lines.append(f"- {takeaway}")
        return {"status": "success", "markdown": "\n".join(lines)}
    except Exception as e:
        return {"status": "error", "markdown": f"Failed to format explanation: {e}"}


# ─────────────────────────────────────────────────────────────────────────────
# Export Tools
# ─────────────────────────────────────────────────────────────────────────────


def _default_export_dir() -> str:
    """Return the user's Documents/Study-Agent Documents folder."""
    target_dir = os.path.join(os.path.expanduser("~"), "Documents", "Study-Agent Documents")
    os.makedirs(target_dir, exist_ok=True)
    return target_dir


def _clean_for_pdf(text: str) -> str:
    """Replace emojis and non-Latin-1 characters with safe text/symbols for standard FPDF fonts."""
    replacements = {
        "📇": "[Flashcards]",
        "📚": "[Study Guide]",
        "📝": "[Quiz]",
        "🧪": "[Practice Test]",
        "💡": "Tip:",
        "📌": "Summary:",
        "👶": "Simplified:",
        "🎭": "Analogy:",
        "🗝️": "Key Takeaways:",
        "✅": "[Answer Key]",
        "☐": "[ ]",
        "\u2022": "-",
    }
    for char, rep in replacements.items():
        text = text.replace(char, rep)
    # Standard FPDF fonts only support Latin-1. Encode and decode to replace other Unicode characters.
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _parse_markdown_lines(markdown_content: str) -> list[dict]:
    """Parse Markdown text into a list of typed token dicts for rendering.

    Supported tokens:
      {"type": "h1"|"h2"|"h3", "text": str}
      {"type": "hr"}
      {"type": "blockquote", "text": str}
      {"type": "bullet", "text": str, "indent": int}
      {"type": "paragraph", "text": str, "bold_spans": list[tuple[int,int]]}
      {"type": "blank"}
    """
    import re

    tokens = []
    for raw in markdown_content.splitlines():
        line = raw.rstrip()

        if re.match(r"^#{3}\s", line):
            tokens.append({"type": "h3", "text": line[4:].strip()})
        elif re.match(r"^#{2}\s", line):
            tokens.append({"type": "h2", "text": line[3:].strip()})
        elif re.match(r"^#\s", line):
            tokens.append({"type": "h1", "text": line[2:].strip()})
        elif re.match(r"^---+$", line):
            tokens.append({"type": "hr"})
        elif re.match(r"^>{1}\s?", line):
            tokens.append({"type": "blockquote", "text": re.sub(r"^>\s?", "", line)})
        elif re.match(r"^(\s*[-*])\s", line):
            indent = len(line) - len(line.lstrip())
            text = re.sub(r"^\s*[-*]\s", "", line)
            tokens.append({"type": "bullet", "text": text, "indent": indent})
        elif line.strip() == "":
            tokens.append({"type": "blank"})
        else:
            # Find bold spans (**text**)
            bold_spans = []
            clean = ""
            pos = 0
            for m in re.finditer(r"\*\*(.+?)\*\*", line):
                clean += line[pos : m.start()] + m.group(1)
                bold_spans.append((len(clean) - len(m.group(1)), len(clean)))
                pos = m.end()
            clean += line[pos:]
            # Strip remaining single-star italics for plain text
            clean = re.sub(r"_(.+?)_", r"\1", clean)
            tokens.append(
                {"type": "paragraph", "text": clean, "bold_spans": bold_spans}
            )

    return tokens


def export_as_pdf(markdown_content: str, filename: str, output_dir: str) -> dict:
    """Export formatted Markdown study material as a PDF file.

    Args:
        markdown_content: The full Markdown text of the study material to export.
        filename: Desired filename WITHOUT extension (e.g. "photosynthesis_flashcards").
        output_dir: Directory path where the file should be saved. Use "default" to save to ~/Documents/Study-Agent Documents.

    Returns:
        A dict with 'status' and 'file_path' keys.
    """
    try:
        from fpdf import FPDF  # type: ignore
    except ImportError:
        return {
            "status": "error",
            "file_path": "fpdf2 not installed. Run: uv add fpdf2",
        }

    save_dir = (
        _default_export_dir() if output_dir.strip().lower() == "default" else output_dir
    )
    os.makedirs(save_dir, exist_ok=True)
    safe_name = "".join(
        c if c.isalnum() or c in "-_ " else "_" for c in filename
    ).strip()
    file_path = os.path.join(save_dir, f"{safe_name}.pdf")

    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()
        pdf.set_margins(20, 20, 20)

        # Page width available for text
        W = pdf.w - 40

        tokens = _parse_markdown_lines(markdown_content)
        for tok in tokens:
            t = tok["type"]

            if t == "h1":
                pdf.set_font("Helvetica", "B", 20)
                pdf.set_text_color(30, 30, 80)
                pdf.multi_cell(W, 10, _clean_for_pdf(tok["text"]), ln=True)
                pdf.set_draw_color(80, 80, 200)
                pdf.set_line_width(0.5)
                pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
                pdf.ln(4)

            elif t == "h2":
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 14)
                pdf.set_text_color(40, 40, 120)
                pdf.multi_cell(W, 8, _clean_for_pdf(tok["text"]), ln=True)
                pdf.ln(1)

            elif t == "h3":
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(60, 60, 150)
                pdf.multi_cell(W, 7, _clean_for_pdf(tok["text"]), ln=True)

            elif t == "hr":
                pdf.ln(2)
                pdf.set_draw_color(180, 180, 200)
                pdf.set_line_width(0.3)
                pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
                pdf.ln(3)

            elif t == "blockquote":
                pdf.set_font("Helvetica", "I", 10)
                pdf.set_text_color(80, 80, 100)
                pdf.set_x(28)
                pdf.set_draw_color(150, 150, 200)
                pdf.set_line_width(0.8)
                y = pdf.get_y()
                pdf.multi_cell(W - 8, 6, _clean_for_pdf(tok["text"]), ln=True)
                pdf.line(22, y, 22, pdf.get_y())
                pdf.ln(1)

            elif t == "bullet":
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 40)
                indent = 8 + tok.get("indent", 0) * 3
                pdf.set_x(20 + indent)
                bullet_text = f"-  {tok['text']}"
                pdf.multi_cell(W - indent, 6, _clean_for_pdf(bullet_text), ln=True)

            elif t == "blank":
                pdf.ln(2)

            elif t == "paragraph":
                text = tok["text"]
                bold_spans = tok.get("bold_spans", [])
                if bold_spans:
                    # Render mixed bold/normal text using write()
                    pdf.set_text_color(40, 40, 40)
                    pos = 0
                    for start, end in bold_spans:
                        if pos < start:
                            pdf.set_font("Helvetica", "", 10)
                            pdf.write(6, _clean_for_pdf(text[pos:start]))
                        pdf.set_font("Helvetica", "B", 10)
                        pdf.write(6, _clean_for_pdf(text[start:end]))
                        pos = end
                    if pos < len(text):
                        pdf.set_font("Helvetica", "", 10)
                        pdf.write(6, _clean_for_pdf(text[pos:]))
                    pdf.ln(7)
                else:
                    pdf.set_font("Helvetica", "", 10)
                    pdf.set_text_color(40, 40, 40)
                    pdf.multi_cell(W, 6, _clean_for_pdf(text), ln=True)

        pdf.output(file_path)
        return {"status": "success", "file_path": file_path}

    except Exception as e:
        return {"status": "error", "file_path": f"PDF export failed: {e}"}


def export_as_docx(markdown_content: str, filename: str, output_dir: str) -> dict:
    """Export formatted Markdown study material as a Microsoft Word (.docx) file.

    Args:
        markdown_content: The full Markdown text of the study material to export.
        filename: Desired filename WITHOUT extension (e.g. "photosynthesis_study_guide").
        output_dir: Directory path where the file should be saved. Use "default" to save to ~/Documents/Study-Agent Documents.

    Returns:
        A dict with 'status' and 'file_path' keys.
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import RGBColor  # type: ignore
    except ImportError:
        return {
            "status": "error",
            "file_path": "python-docx not installed. Run: uv add python-docx",
        }

    save_dir = (
        _default_export_dir() if output_dir.strip().lower() == "default" else output_dir
    )
    os.makedirs(save_dir, exist_ok=True)
    safe_name = "".join(
        c if c.isalnum() or c in "-_ " else "_" for c in filename
    ).strip()
    file_path = os.path.join(save_dir, f"{safe_name}.docx")

    try:
        doc = Document()

        # Set narrow-ish margins (1 inch = 914400 EMUs / 914400 * 1.2 = ~1.2 in)
        for section in doc.sections:
            section.top_margin = section.bottom_margin = 914400  # 1 inch
            section.left_margin = section.right_margin = 1143000  # 1.25 inch

        def _add_bold_run(para, text: str, bold_spans: list) -> None:
            """Add paragraph runs with selective bolding based on char spans."""
            if not bold_spans:
                para.add_run(text)
                return
            pos = 0
            for start, end in bold_spans:
                if pos < start:
                    para.add_run(text[pos:start])
                run = para.add_run(text[start:end])
                run.bold = True
                pos = end
            if pos < len(text):
                para.add_run(text[pos:])

        tokens = _parse_markdown_lines(markdown_content)
        for tok in tokens:
            t = tok["type"]

            if t == "h1":
                p = doc.add_heading(tok["text"], level=1)
                p.runs[0].font.color.rgb = RGBColor(30, 30, 80)

            elif t == "h2":
                p = doc.add_heading(tok["text"], level=2)
                p.runs[0].font.color.rgb = RGBColor(40, 40, 120)

            elif t == "h3":
                p = doc.add_heading(tok["text"], level=3)
                p.runs[0].font.color.rgb = RGBColor(60, 60, 150)

            elif t == "hr":
                # Horizontal rule via bottom border on an empty paragraph
                p = doc.add_paragraph()
                from docx.oxml import OxmlElement  # type: ignore
                from docx.oxml.ns import qn  # type: ignore

                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "AAAACC")
                pBdr.append(bottom)
                pPr.append(pBdr)

            elif t == "blockquote":
                p = doc.add_paragraph(style="Quote")
                p.add_run(tok["text"]).italic = True

            elif t == "bullet":
                indent = tok.get("indent", 0)
                style = "List Bullet 2" if indent > 0 else "List Bullet"
                p = doc.add_paragraph(style=style)
                p.add_run(tok["text"])

            elif t == "blank":
                # Small spacing — use empty paragraph sparingly
                pass  # paragraph breaks handled by heading/paragraph spacing

            elif t == "paragraph":
                p = doc.add_paragraph()
                _add_bold_run(p, tok["text"], tok.get("bold_spans", []))

        doc.save(file_path)
        return {"status": "success", "file_path": file_path}

    except Exception as e:
        return {"status": "error", "file_path": f"DOCX export failed: {e}"}
